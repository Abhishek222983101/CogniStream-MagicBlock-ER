"""
Ingest Router - POST /api/ingest
Accepts unstructured patient data (PDF, DOCX, raw clinical text)
and uses LLM to extract structured patient JSON.

This is the "semi-structured / unstructured data" handler
required by the problem statement.
"""

import io
import json
import logging
import re
from typing import Optional

from fastapi import APIRouter, HTTPException, UploadFile, File, Form

router = APIRouter()
logger = logging.getLogger(__name__)

# ─── System prompt for Mistral to extract structured patient data ────────────

EXTRACT_SYSTEM_PROMPT = """You are a clinical data extraction AI. Given raw, messy, or semi-structured patient medical records, extract a structured JSON patient record.

You MUST return ONLY valid JSON with this exact schema (fill in what you can find, use null for missing fields):

{
  "patient_id": "EXTRACTED_001",
  "demographics": {
    "age": <integer or null>,
    "gender": "<Male/Female/Other or null>",
    "name": "<patient name or null>",
    "city": "<city or null>",
    "state": "<state or null>",
    "lat": <latitude float or null>,
    "lng": <longitude float or null>
  },
  "diagnosis": {
    "primary": "<primary diagnosis>",
    "subtype": "<subtype or null>",
    "stage": "<Stage I/II/III/IV or null>",
    "icd10": "<ICD-10 code or null>",
    "biomarkers": <{"marker": "value"} object or null>
  },
  "medical_history": ["<comorbidity1>", "<comorbidity2>"],
  "medications": ["<med1>", "<med2>"],
  "prior_treatments": [],
  "lab_values": {
    "hemoglobin": <float or null>,
    "wbc": <float or null>,
    "platelets": <float or null>,
    "creatinine": <float or null>,
    "bilirubin": <float or null>,
    "alt": <float or null>,
    "ast": <float or null>
  },
  "ecog_status": <0-4 integer or null>,
  "smoking_status": "<Never/Former/Current or null>",
  "allergies": ["<allergy1>"],
  "clinical_notes": "<1-2 sentence summary of the case>"
}

Important rules:
- Extract ALL medical information you can find
- For Indian cities, try to include lat/lng if you know them
- If age is not directly stated but DOB is given, calculate age
- Convert all lab values to standard units
- Infer ECOG status from functional descriptions if not explicitly stated
- Return ONLY the JSON object, no explanation text"""

EXTRACT_USER_PROMPT = """Extract a structured patient record from this clinical text:

---
{text}
---

Return ONLY valid JSON matching the schema."""


def _extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF using PyPDF2."""
    try:
        import PyPDF2

        reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
        text_parts = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        return "\n".join(text_parts)
    except Exception as e:
        logger.error(f"PDF extraction failed: {e}")
        raise HTTPException(status_code=400, detail=f"Could not read PDF: {str(e)}")


def _extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX using python-docx."""
    try:
        import docx

        doc = docx.Document(io.BytesIO(file_bytes))
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    text_parts.append(row_text)
        return "\n".join(text_parts)
    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        raise HTTPException(status_code=400, detail=f"Could not read DOCX: {str(e)}")


def _try_parse_json(text: str) -> Optional[dict]:
    """Try to extract a JSON object from LLM response text."""
    # Direct parse
    try:
        return json.loads(text.strip())
    except json.JSONDecodeError:
        pass

    # Find JSON block in markdown code fence
    match = re.search(r"```(?:json)?\s*(\{.*?\})\s*```", text, re.DOTALL)
    if match:
        try:
            return json.loads(match.group(1))
        except json.JSONDecodeError:
            pass

    # Find any JSON object
    match = re.search(r"\{[\s\S]*\}", text)
    if match:
        try:
            return json.loads(match.group())
        except json.JSONDecodeError:
            pass

    return None


@router.post("/api/ingest")
async def ingest_patient_record(
    file: Optional[UploadFile] = File(None),
    clinical_text: Optional[str] = Form(None),
):
    """
    Ingest an unstructured or semi-structured patient record.

    Accepts:
    - PDF file upload (.pdf)
    - DOCX file upload (.docx)
    - JSON file upload (.json) — passed through directly
    - Raw clinical text (form field)

    Returns:
    - extracted_text: the raw text extracted from the file
    - structured_data: the LLM-parsed patient JSON
    - source_type: "pdf" | "docx" | "json" | "text"
    - confidence: how confident the extraction was
    """
    from backend.main import get_mistral_client

    extracted_text = ""
    source_type = "text"

    # ─── Handle file upload ──────────────────────────────────────────────────
    if file is not None:
        file_bytes = await file.read()
        filename = file.filename or ""
        ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""

        if ext == "json":
            # JSON — try to parse directly, no LLM needed
            try:
                data = json.loads(file_bytes.decode("utf-8"))
                return {
                    "extracted_text": json.dumps(data, indent=2)[:2000],
                    "structured_data": data,
                    "source_type": "json",
                    "confidence": 100,
                    "llm_used": False,
                    "extraction_notes": "JSON file parsed directly — no LLM extraction needed.",
                }
            except json.JSONDecodeError as e:
                raise HTTPException(
                    status_code=400, detail=f"Invalid JSON file: {str(e)}"
                )

        elif ext == "pdf":
            extracted_text = _extract_text_from_pdf(file_bytes)
            source_type = "pdf"

        elif ext in ("docx", "doc"):
            extracted_text = _extract_text_from_docx(file_bytes)
            source_type = "docx"

        elif ext == "txt":
            extracted_text = file_bytes.decode("utf-8", errors="replace")
            source_type = "text"

        else:
            raise HTTPException(
                status_code=400,
                detail=f"Unsupported file type '.{ext}'. Accepted: .json, .pdf, .docx, .txt",
            )

    elif clinical_text is not None and clinical_text.strip():
        extracted_text = clinical_text.strip()
        source_type = "text"

    else:
        raise HTTPException(
            status_code=400,
            detail="No input provided. Upload a file or provide clinical_text.",
        )

    # ─── Validate we got text ────────────────────────────────────────────────
    if not extracted_text.strip():
        raise HTTPException(
            status_code=400, detail="No text could be extracted from the file."
        )

    logger.info(f"Ingested {source_type} input: {len(extracted_text)} chars")

    # ─── Use Mistral API to parse unstructured text into structured JSON ─────
    mistral = get_mistral_client()
    if mistral is None:
        raise HTTPException(
            status_code=503,
            detail="Mistral API client not available. Cannot parse unstructured text.",
        )

    try:
        # Truncate very long texts to avoid token limits
        text_for_llm = extracted_text[:8000]

        response_text = await mistral.chat(
            user_message=EXTRACT_USER_PROMPT.format(text=text_for_llm),
            system_message=EXTRACT_SYSTEM_PROMPT,
            temperature=0.1,
            max_tokens=2048,
        )

        structured_data = _try_parse_json(response_text)

        if structured_data is None:
            return {
                "extracted_text": extracted_text[:2000],
                "structured_data": None,
                "source_type": source_type,
                "confidence": 0,
                "llm_used": True,
                "extraction_notes": "LLM could not produce valid JSON. Raw response included.",
                "llm_raw_response": response_text[:2000],
            }

        # Ensure patient_id exists
        if not structured_data.get("patient_id"):
            structured_data["patient_id"] = (
                f"INGEST_{source_type.upper()}_{hash(extracted_text[:100]) % 10000:04d}"
            )

        return {
            "extracted_text": extracted_text[:2000],
            "structured_data": structured_data,
            "source_type": source_type,
            "confidence": 85,  # LLM extraction confidence
            "llm_used": True,
            "extraction_notes": f"Successfully extracted structured patient data from {source_type.upper()} using Mistral LLM.",
        }

    except Exception as e:
        logger.error(f"LLM extraction failed: {e}")
        raise HTTPException(status_code=500, detail=f"LLM extraction failed: {str(e)}")
